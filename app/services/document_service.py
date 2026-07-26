"""
services/document_service.py
------------------------------
Extracts raw text from an uploaded complaint document so it can be fed
into the LangGraph `parse_input_node`. Supports the formats the frontend
advertises: PDF, DOCX, TXT, EML.

Each format has its own extractor function; `extract_text()` is the single
entry point routers should call — it picks the right extractor from the
filename's extension and raises a clear error for anything else.
"""

import io
from email import policy
from email.parser import BytesParser

import pdfplumber
from docx import Document

SUPPORTED_EXTENSIONS = (".pdf", ".docx", ".txt", ".eml")


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract and concatenate text from every page of a PDF's bytes."""
    text_chunks = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            if page_text.strip():
                text_chunks.append(page_text)

    if not text_chunks:
        raise ValueError("No extractable text found in PDF (it may be a scanned image).")

    return "\n".join(text_chunks)


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract paragraph and table text from a Word document."""
    doc = Document(io.BytesIO(file_bytes))
    text_chunks = [p.text for p in doc.paragraphs if p.text.strip()]

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_chunks.append(row_text)

    if not text_chunks:
        raise ValueError("No extractable text found in the DOCX file.")

    return "\n".join(text_chunks)


def extract_text_from_txt(file_bytes: bytes) -> str:
    """Decode a plain-text file, tolerating non-UTF-8 encodings."""
    for encoding in ("utf-8", "latin-1"):
        try:
            text = file_bytes.decode(encoding)
            if text.strip():
                return text
            raise ValueError("The TXT file is empty.")
        except UnicodeDecodeError:
            continue
    raise ValueError("Could not decode the TXT file as text.")


def extract_text_from_eml(file_bytes: bytes) -> str:
    """Extract sender, subject, and body text from a .eml email file."""
    msg = BytesParser(policy=policy.default).parsebytes(file_bytes)

    header_lines = []
    if msg.get("From"):
        header_lines.append(f"From: {msg['From']}")
    if msg.get("Subject"):
        header_lines.append(f"Subject: {msg['Subject']}")
    if msg.get("Date"):
        header_lines.append(f"Date: {msg['Date']}")

    body = ""
    if msg.is_multipart():
        for part in msg.walk():
            if part.get_content_type() == "text/plain":
                body = part.get_content()
                break
        if not body:
            for part in msg.walk():
                if part.get_content_type() == "text/html":
                    body = part.get_content()
                    break
    else:
        body = msg.get_content()

    combined = "\n".join(header_lines) + "\n\n" + (body or "")
    if not combined.strip():
        raise ValueError("No extractable text found in the EML file.")

    return combined


def extract_text(filename: str, file_bytes: bytes) -> str:
    """Dispatch to the right extractor based on the file's extension."""
    lower_name = filename.lower()

    if lower_name.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    if lower_name.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    if lower_name.endswith(".txt"):
        return extract_text_from_txt(file_bytes)
    if lower_name.endswith(".eml"):
        return extract_text_from_eml(file_bytes)

    raise ValueError(
        f"Unsupported file type. Supported formats: {', '.join(SUPPORTED_EXTENSIONS)}."
    )